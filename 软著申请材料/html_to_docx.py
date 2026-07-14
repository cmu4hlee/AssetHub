import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def convert_source_code_html_to_docx(html_path, output_path, software_name, version):
    """将源代码 HTML 转换为 Word 文档（专为代码设计）"""

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Courier New'
    style.font.size = Pt(9)

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # 添加标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{software_name} 源代码")
    run.font.size = Pt(18)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"版本：{version}")
    run.font.size = Pt(12)

    doc.add_paragraph()

    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'lxml')
    pages = soup.find_all('div', class_='page')

    for page_idx, page in enumerate(pages, 1):
        code_div = page.find('div', class_='code')
        footer = page.find('div', class_='footer')

        if page_idx > 1:
            doc.add_page_break()

        # 添加页眉
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{software_name} 源代码 - {version}")
        run.font.size = Pt(9)

        if code_div:
            lines = code_div.find_all('span', class_='line-num')
            for line in lines:
                line_num = line.get_text(strip=True)
                code_text = line.next_sibling
                if code_text:
                    code_text = str(code_text)
                else:
                    code_text = ""

                # 处理 HTML 实体
                code_text = code_text.replace('&lt;', '<').replace('&gt;', '>')
                code_text = code_text.replace('&amp;', '&').replace('&quot;', '"')
                code_text = code_text.replace('&#039;', "'").replace('&nbsp;', ' ')

                p = doc.add_paragraph()
                # 行号
                run = p.add_run(f"{line_num:>4}  ")
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                run.font.color.rgb = None

                # 代码内容
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)

                # 设置段前段后间距以节省空间
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0

        # 添加页脚
        if footer:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(footer.get_text(strip=True))
            run.font.size = Pt(9)

    doc.save(output_path)
    print(f'已生成 Word 文档：{output_path}')
    print(f'总页数：{len(pages)}')


if __name__ == '__main__':
    base_dir = '/Volumes/移动硬盘（500）/AssetHub/软著申请材料'
    html_path = os.path.join(base_dir, '01-源代码.html')
    output_path = os.path.join(base_dir, '01-源代码.docx')

    software_name = 'AssetHost智能资产管理平台'
    version = 'V1.0'

    convert_source_code_html_to_docx(html_path, output_path, software_name, version)
