import sys
import os
import re
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = tcPr.find(qn(tag))
                if element is None:
                    element = docx.oxml.OxmlElement(tag)
                    tcPr.append(element)
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), edge_data.get('color', '000000'))


def set_chinese_font(run, font_name='宋体', size=10.5, bold=False):
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(size)
    font.bold = bold
    if font_name == '黑体':
        font.color.rgb = RGBColor(0, 0, 0)


def convert_html_to_docx(html_content, output_path, title=None, base_dir=None):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    soup = BeautifulSoup(html_content, 'lxml')

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        set_chinese_font(run, font_name='黑体', size=18, bold=True)
        doc.add_paragraph()

    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'pre', 'blockquote', 'hr', 'img']):
        if element.name == 'h1':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text(strip=True))
            set_chinese_font(run, font_name='黑体', size=16, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif element.name == 'h2':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text(strip=True))
            set_chinese_font(run, font_name='黑体', size=14, bold=True)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(5)
        elif element.name == 'h3':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text(strip=True))
            set_chinese_font(run, font_name='黑体', size=12, bold=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif element.name == 'h4':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text(strip=True))
            set_chinese_font(run, font_name='黑体', size=11, bold=True)
        elif element.name == 'img':
            src = element.get('src', '')
            if src and base_dir:
                img_path = os.path.join(base_dir, src)
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        run = p.add_run()
                        run.add_picture(img_path, width=Cm(15))
                    except Exception as e:
                        print(f'添加图片失败: {img_path}, 错误: {e}')
            elif src.startswith('screenshot:'):
                screenshot_name = src.replace('screenshot:', '')
                img_path = os.path.join(base_dir, screenshot_name)
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        run = p.add_run()
                        run.add_picture(img_path, width=Cm(15))
                    except Exception as e:
                        print(f'添加图片失败: {img_path}, 错误: {e}')
        elif element.name == 'p':
            text = element.get_text(strip=True)
            if not text:
                continue
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == 'code':
                    run = p.add_run(child.get_text())
                    set_chinese_font(run, font_name='Courier New', size=10)
                    run.font.name = 'Courier New'
                elif child.name == 'strong':
                    run = p.add_run(child.get_text())
                    set_chinese_font(run, font_name='宋体', size=10.5, bold=True)
                elif child.name == 'em':
                    run = p.add_run(child.get_text())
                    set_chinese_font(run, font_name='宋体', size=10.5)
                    run.font.italic = True
                elif child.name == 'a':
                    run = p.add_run(child.get_text())
                    set_chinese_font(run, font_name='宋体', size=10.5)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True
                elif isinstance(child, str):
                    run = p.add_run(child)
                    set_chinese_font(run, font_name='宋体', size=10.5)
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.line_spacing = 1.15
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(li.get_text(strip=True))
                set_chinese_font(run, font_name='宋体', size=10.5)
        elif element.name == 'ol':
            for idx, li in enumerate(element.find_all('li', recursive=False), 1):
                p = doc.add_paragraph(style='List Number')
                run = p.add_run(li.get_text(strip=True))
                set_chinese_font(run, font_name='宋体', size=10.5)
        elif element.name == 'table':
            rows = element.find_all('tr')
            if not rows:
                continue
            col_count = max(len(row.find_all(['td', 'th'])) for row in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    if j >= col_count:
                        break
                    doc_cell = table.rows[i].cells[j]
                    doc_cell.text = cell.get_text(strip=True)
                    for paragraph in doc_cell.paragraphs:
                        for run in paragraph.runs:
                            set_chinese_font(run, font_name='宋体', size=10)
                            if cell.name == 'th':
                                run.font.bold = True
        elif element.name == 'pre':
            code = element.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code)
            set_chinese_font(run, font_name='Courier New', size=9)
        elif element.name == 'blockquote':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text(strip=True))
            set_chinese_font(run, font_name='宋体', size=10.5)
            run.font.italic = True
            p.paragraph_format.left_indent = Inches(0.3)
        elif element.name == 'hr':
            doc.add_paragraph('─' * 50)

    doc.save(output_path)
    print(f'已生成 Word 文档：{output_path}')


def convert_markdown_to_docx(md_path, output_path, title=None):
    base_dir = os.path.dirname(md_path)

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    html_content = markdown.markdown(
        md_content,
        extensions=['tables', 'fenced_code', 'toc', 'sane_lists']
    )

    convert_html_to_docx(html_content, output_path, title, base_dir)


if __name__ == '__main__':
    base_dir = '/Volumes/移动硬盘（500）/AssetHub/软著申请材料'

    files = [
        ('00-材料说明.md', '00-材料说明.docx', 'AssetHost 智能资产管理平台 软著申请材料说明'),
        ('02-软件说明书.md', '02-软件说明书.docx', 'AssetHost 智能资产管理平台 V1.0 软件说明书'),
        ('03-登记申请表信息.md', '03-登记申请表信息.docx', 'AssetHost 智能资产管理平台 登记申请表信息'),
    ]

    for md_file, docx_file, title in files:
        md_path = os.path.join(base_dir, md_file)
        output_path = os.path.join(base_dir, docx_file)
        if os.path.exists(md_path):
            convert_markdown_to_docx(md_path, output_path, title)
        else:
            print(f'文件不存在：{md_path}')
